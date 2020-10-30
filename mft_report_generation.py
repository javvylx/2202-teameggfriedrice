import struct, os, sys, getopt, datetime, csv, hashlib, argparse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date
from docx.oxml import OxmlElement, ns
from docx2pdf import convert

# * Timezone of his script follows host's Timezone
TIMEZONE = datetime.datetime.now().astimezone().tzinfo


# * GET MFT ATTRIBUTE DATA
def get_mft_data(raw_file, mft_record, read_ptr):
    # Generally each file MFT is 1024 btyes
    while read_ptr < 1024:

        attr_record = get_attr_header(raw_file[read_ptr:])
        if attr_record['type'] == 0xffffffff:  # End of attributes
            return mft_record

        if attr_record['name_len'] > 0:
            file_bytes = raw_file[
                         read_ptr + attr_record['name_off']: read_ptr + attr_record['name_off'] + attr_record[
                             'name_len'] * 2]
            attr_record['name'] = file_bytes.decode('utf-16').encode('utf-8')
        else:
            attr_record['name'] = ''

        if attr_record['type'] == 0x10:  # Standard Information
            si_record = get_si_attribute(raw_file[read_ptr + attr_record['soff']:], TIMEZONE)
            mft_record['si'] = si_record

        if attr_record['type'] == 0x30:  # File Name
            fn_record = get_fn_attribute(raw_file[read_ptr + attr_record['soff']:], TIMEZONE)
            mft_record['fncount'] += 1
            mft_record['fn', mft_record['fncount']] = fn_record

        if attr_record['len'] > 0:
            read_ptr = read_ptr + attr_record['len']
        else:
            return mft_record

    return mft_record


# * GET STANDARD INFORMATION ATTRIBUTES
def get_si_attribute(s, tz):
    si_attr = {
        'crtime': get_time(struct.unpack("<L", s[:4])[0], struct.unpack("<L", s[4:8])[0], tz),
        # File create time; "<": little endian, "L": unassigned long (4 bytes)
        'mtime': get_time(struct.unpack("<L", s[8:12])[0], struct.unpack("<L", s[12:16])[0], tz),
        # MFT changed time; "<": little endian, "L": unassigned long (4 bytes)
        'ctime': get_time(struct.unpack("<L", s[16:20])[0], struct.unpack("<L", s[20:24])[0], tz),
        # Entry time; "<": little endian, "L": unassigned long (4 bytes)
        'atime': get_time(struct.unpack("<L", s[24:28])[0], struct.unpack("<L", s[28:32])[0], tz),
        # File alter time; "<": little endian, "L": unassigned long (4 bytes)
        'dos': struct.unpack("<I", s[32:36])[0], 'maxver': struct.unpack("<I", s[36:40])[0],
        'ver': struct.unpack("<I", s[40:44])[0], 'class_id': struct.unpack("<I", s[44:48])[0],
        'own_id': struct.unpack("<I", s[48:52])[0], 'sec_id': struct.unpack("<I", s[52:56])[0],
        'quota': struct.unpack("<d", s[56:64])[0], 'usn': struct.unpack("<q", s[64:72])[0],
    }

    return si_attr


# * GET FILE NAME ATTRIBUTES
def get_fn_attribute(s, tz):
    fn_attr = {
        'par_ref': struct.unpack("<Lxx", s[:6])[0], 'par_seq': struct.unpack("<H", s[6:8])[0],
        'crtime': get_time(struct.unpack("<L", s[8:12])[0], struct.unpack("<L", s[12:16])[0], tz),
        'mtime': get_time(struct.unpack("<L", s[16:20])[0], struct.unpack("<L", s[20:24])[0], tz),
        'ctime': get_time(struct.unpack("<L", s[24:28])[0], struct.unpack("<L", s[28:32])[0], tz),
        'atime': get_time(struct.unpack("<L", s[32:36])[0], struct.unpack("<L", s[36:40])[0], tz),
        'alloc_fsize': struct.unpack("<q", s[40:48])[0],
        'real_fsize': struct.unpack("<q", s[48:56])[0],
        'flags': struct.unpack("<q", s[56:64])[0],
        'nlen': struct.unpack("B", bytes([s[64]]))[0],
        'nspace': struct.unpack("B", bytes([s[65]]))[0],
    }

    attr_bytes = s[66:66 + fn_attr['nlen'] * 2]
    try:
        fn_attr['name'] = attr_bytes.decode('utf-16').encode('utf-8')
    except:
        fn_attr['name'] = 'UnableToDecodeFilename'

    return fn_attr


# * CONVERT FILETIME TO UNIX TIME
def get_time(low, high, tz):
    if low == 0 and high == 0:
        dt = 0
        dt_str = 'Undefined'
        unix_time = 0
        return
    unix_time = (float(high) * 2 ** 32 + low) * 1e-7 - 11644473600

    try:

        if tz:
            dt = datetime.datetime.fromtimestamp(unix_time)
        else:
            dt = datetime.datetime.utcfromtimestamp(unix_time)

        # Edit timestamp delimiter
        dt_str = dt.isoformat(' ')
    except:
        dt = 0
        dt_str = 'Invalid timestamp'
        unix_time = 0
    return dt, dt_str, unix_time


# * CHECK FOR MFT RECORD ANOMALIES
def anomaly_mftrecord_check(mft_record):
    if mft_record['end_seq_1'] != mft_record['fixup_data'] or mft_record['end_seq_2'] != mft_record['fixup_data']:
        return 1
    return 0


# * CHECK FOR $SI TIMESTAMP ANOMALIES
def anomaly_timestamp_check_SI(mft_record):
    '''
    Flags for $SI Checks
    0: Okay
    1. 0s in $SI
    2: Invalid $SI Timestamp
    3. $SI nanosec = 0
    4. $SI (crtime) after $SI (mtime)
    5. $SI (atime) after $SI (mtime) and $SI (atime) after $SI (crtime)
    6. $SI after Current Time
    '''

    # Check for any only 0s in $SI
    if mft_record['si']['crtime'] is None or mft_record['si']['mtime'] is None or mft_record['si']['ctime'] is None or \
            mft_record['si']['atime'] is None:
        return 1

    # Check for any invalid timestamp in $SI
    elif mft_record['si']['crtime'][0] == 0 or mft_record['si']['mtime'][0] == 0 or mft_record['si']['ctime'][0] == 0 or \
            mft_record['si']['atime'][0] == 0:
        return 2

    # Check for $SI with a nanosecond value of '0'
    else:
        if mft_record['si']['crtime'][2] != 0:
            if len(str(mft_record['si']['crtime'][2]).split('.')[1]) == 1 and str(mft_record['si']['crtime'][2]).split('.')[1] == str(0):
                return 3
            elif str(mft_record['si']['crtime'][2]).split('.')[1][-3] == str(000):
                return 3
        elif mft_record['si']['mtime'][2] != 0:
            if len(str(mft_record['si']['mtime'][2]).split('.')[1]) == 1 and str(mft_record['si']['mtime'][2]).split('.')[1] == str(0):
                return 3
            elif str(mft_record['si']['mtime'][2]).split('.')[1][-3] == str(000):
                return 3
        elif mft_record['si']['ctime'][2] != 0:
            if len(str(mft_record['si']['ctime'][2]).split('.')[1]) == 1 and str(mft_record['si']['ctime'][2]).split('.')[1] == str(0):
                return 3
            elif str(mft_record['si']['ctime'][2]).split('.')[1][-3] == str(000):
                return 3
        elif mft_record['si']['atime'][2] != 0:
            if len(str(mft_record['si']['atime'][2]).split('.')[1]) == 1 and str(mft_record['si']['atime'][2]).split('.')[1] == str(0):
                return 3
            elif str(mft_record['si']['atime'][2]).split('.')[1][-3] == str(000):
                return 3

        # Check for $SI Create Time after the $SI Modify Time
        if mft_record['si']['crtime'][1] > mft_record['si']['mtime'][1]:
            return 4

        # Check for #SI Access Time after the $SI Modify and $SI Create Time
        if mft_record['si']['atime'][1] > mft_record['si']['mtime'][1] and mft_record['si']['atime'][1] > \
                mft_record['si']['crtime'][1]:
            return 5

        # Check for $SI Time after Current Time
        if mft_record['si']['crtime'][1] > str(datetime.datetime.now()) or mft_record['si']['mtime'][1] > str(
                datetime.datetime.now()) or mft_record['si']['ctime'][1] > str(datetime.datetime.now()) or \
                mft_record['si']['atime'][1] > str(datetime.datetime.now()):
            return 6

    # Return Flag 0 for no anomalies    
    return 0


# * CHECK FOR $FN & $SI TIMESTAMP ANOMALIES
def anomaly_timestamp_check_SI_FN(mft_record, fn_count):
    '''
    Flags for both $SI and $FN
    0: Okay
    1: 0s in $FN
    2: Invalid $FN Timestamp
    3. $FN nanosec = 0
    4: $SI (crtime) before $FN (crtime)
    5: $FN after Current Time
    '''
    check = 1
    while check <= fn_count:

        # Check for any only 0s in $FN
        if mft_record['fn', check]['crtime'] is None or mft_record['fn', check]['mtime'] is None or \
                mft_record['fn', check]['ctime'] is None or mft_record['fn', check]['atime'] is None:
            return 1

        # Check for any invalid timestamp in $FN
        elif mft_record['fn', check]['crtime'][0] == 0 or mft_record['fn', check]['mtime'][0] == 0 or \
                mft_record['fn', check]['ctime'][0] == 0 or mft_record['fn', check]['atime'][0] == 0:
            return 2

        else:

            if mft_record['fn', check]['crtime'][2] != 0:
                if len(str(mft_record['fn', check]['crtime'][2]).split('.')[1]) == 1 and str(mft_record['fn', check]['crtime'][2]).split('.')[1] == str(0):
                    return 3
                elif str(mft_record['fn', check]['crtime'][2]).split('.')[1][-3] == str(000):
                    return 3
            elif mft_record['fn', check]['mtime'][2] != 0:
                if len(str(mft_record['fn', check]['mtime'][2]).split('.')[1]) == 1 and str(mft_record['fn', check]['mtime'][2]).split('.')[1] == str(0):
                    return 3
                elif str(mft_record['fn', check]['mtime'][2]).split('.')[1][-3] == str(000):
                    return 3
            elif mft_record['fn', check]['ctime'][2] != 0:
                if len(str(mft_record['fn', check]['ctime'][2]).split('.')[1]) == 1 and str(mft_record['fn', check]['ctime'][2]).split('.')[1] == str(0):
                    return 3
                elif str(mft_record['fn', check]['ctime'][2]).split('.')[1][-3] == str(000):
                    return 3
            elif mft_record['fn', check]['atime'][2] != 0:
                if len(str(mft_record['fn', check]['atime'][2]).split('.')[1]) == 1 and str(mft_record['fn', check]['atime'][2]).split('.')[1] == str(0):
                    return 3
                elif str(mft_record['fn', check]['atime'][2]).split('.')[1][-3] == str(000):
                    return 3

            # Check for $SI Create Time is before $FN Create Time
            if mft_record['si']['crtime'][1] < mft_record['fn', check]['crtime'][1]:
                return 4

            # Check for $FN Time after Current Time
            if mft_record['fn', check]['crtime'][1] > str(datetime.datetime.now()) or mft_record['fn', check]['mtime'][
                1] > str(datetime.datetime.now()) or mft_record['fn', check]['ctime'][1] > str(
                    datetime.datetime.now()) or mft_record['fn', check]['atime'][1] > str(datetime.datetime.now()):
                return 5

        check += 1

    # Return Flag 0 for no anomalies
    return 0

# * GET ATTRIBUTE HEADER
def get_attr_header(pointer):
    d = {'type': struct.unpack("<L", pointer[:4])[
        0]}  # Attribute Type; "<": little endian, "L": unassigned long (4 bytes)
    if d['type'] == 0xffffffff:
        return d
    d['len'] = struct.unpack("<L", pointer[4:8])[0]  # Length; "<": little endian, "L": unassigned long (4 bytes)
    d['res_flag'] = struct.unpack("B", pointer[8:9])[0]  # Non-resident flag; "B": byte (1 byte)
    d['name_len'] = struct.unpack("B", pointer[9:10])[0]  # Name length; "B": byte (1 byte)
    d['name_off'] = struct.unpack("<H", pointer[10:12])[0]  # Offset to the name; "H": unsigned short (2 byte)
    d['flags'] = struct.unpack("<H", pointer[12:14])[0]  # Flags; "H": unsigned short (2 byte)
    d['attr_id'] = struct.unpack("<H", pointer[14:16])[0]  # Attribute ID; "H": unsigned short (2 byte)
    if d['res_flag'] == 0:
        d['ssize'] = struct.unpack("<L", pointer[16:20])[0]  # dwLength
        d['soff'] = struct.unpack("<H", pointer[20:22])[0]  # wAttrOffset
        d['idxflag'] = struct.unpack("B", pointer[22:23])[0]  # uchIndexedTag
    else:
        d['start_vcn'] = struct.unpack("<Q", pointer[16:24])[0]  # n64StartVCN
        d['last_vcn'] = struct.unpack("<Q", pointer[24:32])[0]  # n64EndVCN
        d['run_off'] = struct.unpack("<H", pointer[32:34])[0]  # wDataRunOffset (in clusters, from start of partition?)
        d['compsize'] = struct.unpack("<H", pointer[34:36])[0]  # wCompressionSize
        d['allocsize'] = struct.unpack("<Lxxxx", pointer[40:48])[0]  # n64AllocSize
        d['realsize'] = struct.unpack("<Lxxxx", pointer[48:56])[0]  # n64RealSize
        d['streamsize'] = struct.unpack("<Lxxxx", pointer[56:64])[0]  # n64StreamSize

    return d

# * GET MFT ENTRY HEADER
def get_mft_entry_header(raw_data):
    mft_record = {
        'signature': struct.unpack("<I", raw_data[:4])[0],  # Signature; "<": little endian, "I": unsigned int (4 bytes)
        'fixup_offset': struct.unpack("<H", raw_data[4:6])[0],
        # Fix-up offset; "<": little endian, "H": unsigned short (2 bytes)
        'fixup_arraysize': struct.unpack("<H", raw_data[6:8])[0],
        # Fix-up values; "<": little endian, "H": unsigned short (2 bytes)
        'log_seq_num': struct.unpack("<q", raw_data[8:16])[0],
        # LogFile Seq Number; "<": little endian, "q": long long (8 bytes)
        'seq_num': struct.unpack("<H", raw_data[16:18])[0],
        # Seq Number; "<": little endian, "H": unsigned short (2 bytes)
        'ref_count': struct.unpack("<H", raw_data[18:20])[0],
        # Reference Count; "<": little endian, "H": unsigned short (2 bytes)
        'attr_offset': struct.unpack("<H", raw_data[20:22])[0],
        # Attribute offset; "<": little endian, "H": unsigned short (2 bytes)
        'entry_flags': struct.unpack("<H", raw_data[22:24])[0],
        # Entry Flags; "<": little endian, "H": unsigned short (2 bytes)
        'used_entry_size': struct.unpack("<I", raw_data[24:28])[0],
        # used entry size; "<": little endian, "I": unsigned int (4 bytes)
        'total_entry_size': struct.unpack("<I", raw_data[28:32])[0],
        # Total entry size (possible > 1024 bytes); "<": little endian, "I": unsigned int (4 bytes)
        'base_rec_file_ref': struct.unpack("<Lxx", raw_data[32:38])[0],
        # Base record file reference; "<": little endian, "I": unsigned int (4 bytes), "x": just padding only (1 byte per x)
        'base_rec_file_seq': struct.unpack("<H", raw_data[38:40])[0],
        # Base record file seq; "<": little endian, "H": unsigned short (2 bytes)
        'next_attr_id': struct.unpack("<H", raw_data[40:42])[0],
        # Next attribute identifier; "<": little endian, "H": unsigned short (2 bytes)
        'record_num': struct.unpack("<I", raw_data[44:48])[0],
        # MFT Entry record number; "<": little endian, "I": unsigned int (4 bytes)
        'fixup_data': raw_data[struct.unpack("<H", raw_data[4:6])[0]:struct.unpack("<H", raw_data[4:6])[0] + 2],
        'end_seq_1': raw_data[510:512],
        'end_seq_2': raw_data[1022:],
        'fncount': 0  # Number of $FN
    }
    return mft_record

# * HEADERS FOR OUT.CSV
def csvheader():
    listheader = ["Data Number", "Signature", "Fix-up Data Offset", "Fix-up Data Array Size", "Log Sequence Number",
                  "Sequence Number", "Reference Count", "Attribute Offset", "Entry Flags", "MFT Record Actual Size",
                  "MFT Record Allocated",
                  "Base Record File Reference", "Next Attribute ID", "MFT Data Number", "$SI Creation Time",
                  "$SI Modified Time",
                  "$SI Change Time", "$SI Access Time", "$FN_1 Creation Time", "$FN_1 Modified Time",
                  "$FN_1 Change Time", "$FN_1 Access Time", "$FN_1 Filename Length", "$FN_1 Name Space Length",
                  "$FN_1 Filename",
                  "$FN_2 Creation Time", "$FN_2 Modified Time", "$FN_2 Change Time", "$FN_2 Access Time",
                  "$FN_2 Filename Length", "$FN_2 Name Space Length", "$FN_2 Filename",
                  "$FN_3 Creation Time", "$FN_3 Modified Time", "$FN_3 Change Time", "$FN_3 Access Time",
                  "$FN_3 Filename Length", "$FN_3 Name Space Length", "$FN_3 Filename"]
    return listheader

# * EXCEL FRIENDLY DATE FORMAT
def excel_date(date1):
    return '="{}"'.format(date1)

# * DICT OF ANOMALY REASON
def retrReason(n):
    anomreason = {
        "anomalymftrecord" : "MFT record might be tampered/damaged",
        "anomaly_SI_1" : "$SI Timestamp zerorised",
        "anomaly_SI_2" : "$SI Timestamp is invalid",
        "anomaly_SI_3" : "$SI Timestamp with 0 nanosecond",
        "anomaly_SI_4" : "$SI Create Timestamp after $SI Modify Timestamp",
        "anomaly_SI_5" : "$SI Assess Timestamp after both $SI Modify Timestamp and $SI Create Timestamp",
        "anomaly_SI_6" : "$SI Timestamp after Current Time",
        "anomaly_SI_FN_1" : "$FN Timestamp zerorised",
        "anomaly_SI_FN_2" : "$FN Timestamp is invalid",
        "anomaly_SI_FN_3" : "$FN Timestamp with 0 nanosecond",
        "anomaly_SI_FN_4" : "$SI Create Timestamp before $FN Create Timestamp",
        "anomaly_SI_FN_5" : "$FN Timestamp after Current Time"
    }
    return anomreason[n]

# * RETURN A STRING TO SEE IF THE SIGNATURE IS GOOD OR BAD
def signaturechk(n):
    if n == 1162627398:
        return "Good MFT Signature"
    else:
        return "Bad MFT Signature"

# * PADDING FOR EXCEL
def padding(n):
    t = []
    for x in range(n):
        t.append('')
    return t

# * GENERATE A LIST TO WRITE INTO OUT.CSV
def writecsv(mft_record):
    csvcontent = [mft_record['datanumber'], signaturechk(mft_record['signature']), 
                    mft_record['fixup_offset'],mft_record['fixup_arraysize'],
                    mft_record['log_seq_num'],mft_record['seq_num'],
                    mft_record['ref_count'],mft_record['attr_offset'],
                    mft_record['entry_flags'],mft_record['used_entry_size'],
                    mft_record['total_entry_size'],mft_record['base_rec_file_ref'],
                    mft_record['next_attr_id'],mft_record['record_num']]
    #print(mft_record['si']['crtime'])
    if 'si' not in mft_record:
        csvcontent.extend(padding(6))
    elif 'si' in mft_record:
        if mft_record['si']['crtime'] is not None:
            csvcontent.extend([excel_date(mft_record['si']['crtime'][1])])
        else:
            csvcontent.extend(padding(1))
        if mft_record['si']['mtime'] is not None:
            csvcontent.extend([excel_date(mft_record['si']['mtime'][1])])
        else:
            csvcontent.extend(padding(1))
        if mft_record['si']['ctime'] is not None:
            csvcontent.extend([excel_date(mft_record['si']['ctime'][1])])
        else:
            csvcontent.extend(padding(1))
        if mft_record['si']['atime'] is not None:
            csvcontent.extend([excel_date(mft_record['si']['atime'][1])])
        else:
            csvcontent.extend(padding(1))
        csvcontent.extend([mft_record['si']['dos'], mft_record['si']["own_id"]])
    if mft_record['fncount'] == 0:
        csvcontent.extend(padding(21))
    elif mft_record['fncount'] > 0:
        for i in range(3):
            if (i) == mft_record['fncount']-1:
                break
            check = i + 1
            if mft_record['fn', check]['crtime'] is not None:
                csvcontent.extend([excel_date(mft_record['fn', check]['crtime'][1])])
            else:
                csvcontent.extend(padding(1))
            if mft_record['fn', check]['mtime'] is not None:
                csvcontent.extend([excel_date(mft_record['fn', check]['mtime'][1])])
            else:
                csvcontent.extend(padding(1))
            if mft_record['fn', check]['ctime'] is not None:
                csvcontent.extend([excel_date(mft_record['fn', check]['ctime'][1])])
            else:
                csvcontent.extend(padding(1))
            if mft_record['fn', check]['atime'] is not None:
                csvcontent.extend([excel_date(mft_record['fn', check]['atime'][1])])
            else:
                csvcontent.extend(padding(1))
            csvcontent.extend([mft_record['fn', check]['flags'], mft_record['fn',check]['nlen'],mft_record['fn',check]['nspace'], str(mft_record['fn',check]['name'])[2:-1]])
    if mft_record['fncount'] < 3:
        csvcontent.extend(padding((3-mft_record['fncount'])*7))
    return csvcontent

# * MAIN TOOL FUNCTION
def get_mft_eh_val(file):
    result = open("./output/out.csv", "w")
    oa_anom = open("./output/overall_anomalies.csv", "w")
    count = 0
    overalldata, wd = [], []
    listcsvcolumns = csvheader()
    writer = csv.writer(result, delimiter=',')
    wr = csv.writer(oa_anom, delimiter=',')
    writer.writerow(listcsvcolumns)
    anomaly_SI_1, anomaly_SI_2, anomaly_SI_3, anomaly_SI_4, anomaly_SI_5, anomaly_SI_6 = 0, 0, 0, 0, 0, 0
    anomaly_SI_FN_1, anomaly_SI_FN_2, anomaly_SI_FN_3, anomaly_SI_FN_4, anomaly_SI_FN_5 = 0, 0, 0, 0, 0
    anomalymftrecord = 0
    file.seek(0)
    while (True):
        anomaly_bool_not_none, anomaly_bool_not_invalid = True, True
        count += 1
        hexdata = file.read(1024)
        if hexdata[:4] == b'\x00\x00\x00\x00':
            continue
        if hexdata[:4] == b'':
            oa_anom.close()
            result.close()
            break
        if hexdata != "":
            # ? Get MFT entry header
            mft_record = {}
            mft_record = get_mft_entry_header(hexdata)

            mft_record = get_mft_data(hexdata, mft_record, mft_record['attr_offset'])
            mft_record['datanumber'] = count
            x = writecsv(mft_record)
            writer.writerow(x)

            #? Do Anomaly Check if MFT record is tampered
            bool_mftrecord_check = anomaly_mftrecord_check(mft_record)
            if bool_mftrecord_check == 1:
                wd.append((count, "anomalymftrecord"))
                wr.writerow([count, retrReason("anomalymftrecord")])
                anomalymftrecord += 1

            # ? Do Anomaly Check on $SI only
            if 'si' in mft_record:
                flag_SI = anomaly_timestamp_check_SI(mft_record)
                if flag_SI != 0:
                    if flag_SI == 1:
                        wd.append((count, "anomaly_SI_1"))
                        wr.writerow([count, retrReason("anomaly_SI_1")])
                        anomaly_SI_1 += 1
                        anomaly_bool_not_none = False
                    if flag_SI == 2:
                        wd.append((count, "anomaly_SI_2"))
                        wr.writerow([count, retrReason("anomaly_SI_2")])
                        anomaly_SI_2 += 1
                        anomaly_bool_not_invalid = False
                    if flag_SI == 3:
                        wd.append((count, "anomaly_SI_3"))
                        wr.writerow([count, retrReason("anomaly_SI_3")])
                        anomaly_SI_3 += 1
                    if flag_SI == 4:
                        wd.append((count, "anomaly_SI_4"))
                        wr.writerow([count, retrReason("anomaly_SI_4")])
                        anomaly_SI_4 += 1
                    if flag_SI == 5:
                        wd.append((count, "anomaly_SI_5"))
                        wr.writerow([count, retrReason("anomaly_SI_5")])
                        anomaly_SI_5 += 1
                    if flag_SI == 6:
                        wd.append((count, "anomaly_SI_6"))
                        wr.writerow([count, retrReason("anomaly_SI_6")])
                        anomaly_SI_6 += 1

            # ? Do Anomaly check only if both $SI and $FN is available
            if mft_record['fncount'] > 0 and 'si' in mft_record and anomaly_bool_not_none and anomaly_bool_not_invalid:
                flag_SI_FN = anomaly_timestamp_check_SI_FN(mft_record, mft_record['fncount'])
                if flag_SI_FN != 0:
                    if flag_SI_FN == 1:
                        wd.append((count, "anomaly_SI_FN_1"))
                        wr.writerow([count, retrReason("anomaly_SI_FN_1")])
                        anomaly_SI_FN_1 += 1
                    if flag_SI_FN == 2:
                        wd.append((count, "anomaly_SI_FN_2"))
                        wr.writerow([count, retrReason("anomaly_SI_FN_2")])
                        anomaly_SI_FN_2 += 1
                    if flag_SI_FN == 3:
                        wd.append((count, "anomaly_SI_FN_3"))
                        wr.writerow([count, retrReason("anomaly_SI_FN_3")])
                        anomaly_SI_FN_3 += 1
                    if flag_SI_FN == 4:
                        wd.append((count, "anomaly_SI_FN_4"))
                        wr.writerow([count, retrReason("anomaly_SI_FN_4")])
                        anomaly_SI_FN_4 += 1
                    if flag_SI_FN == 5:
                        wd.append((count, "anomaly_SI_FN_5"))
                        wr.writerow([count, retrReason("anomaly_SI_FN_5")])
                        anomaly_SI_FN_5 += 1

        else:
            oa_anom.close()
            result.close()
            break

    overalldata.extend(
        [count, anomalymftrecord, anomaly_SI_1, anomaly_SI_2, anomaly_SI_3, anomaly_SI_4, anomaly_SI_5, anomaly_SI_6,
         anomaly_SI_FN_1, anomaly_SI_FN_2, anomaly_SI_FN_3, anomaly_SI_FN_4, anomaly_SI_FN_5])

# *************************************************Start Report********************************************************#
    file_name = "digital_forensic"
    createReport(count, file, overalldata, wd, file_name)

# * GET SHA 256 HASH FOR THE INPUT FILE
def fileHash(file):
    try:
        f = open(file, "rb")  # Buffered Reader
        hashed = f.read()  # read entire file as bytes
        return hashlib.sha256(hashed).hexdigest()
    except Exception as e:
        return None

# * RETRIEVE THE FILE SIZE
def fileSize(file):
    try:
        return os.path.getsize(file)
    except Exception as e:
        return - 1
        
# * GENERATE THE REPORT
def createReport(count, file, overalldata, wd, fileName):
    document = Document()
    header = document.sections[0].header
    htable = header.add_table(1, 2, Inches(6))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].add_paragraph()
    kh = ht0.add_run()
    kh.add_picture('./images/SIT_logo.png', width=Inches(0.8))
    ht1 = htab_cells[1].add_paragraph('Singapore Institute of Technology')
    ht1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    heading = document.add_heading('Digital Forensic Report', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    reportCreationDate = document.add_paragraph("Report Creation Date: " + str(date.today()))
    reportCreationDate.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# ************************************************ Start of Image Summary ******************************************** #
    document.add_page_break()
    summary = document.add_heading('Image File Summary', 0)
    # print(file.name)
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    document.add_paragraph("The name of the image file is: " + str(os.path.basename(file.name)), style=None)
    document.add_paragraph("The total records in the image file is: " + str(count) + " records", style=None)
    document.add_paragraph("The sha256sum of the image file is: " + str(fileHash(file.name)), style=None)
    document.add_paragraph("The size of the image file is: " + str(fileSize(file.name)) + " bytes", style=None)
# ************************************************ End of Image Summary ********************************************** #
# ************************************************ Start of Anomalies Summary **************************************** #
    # Anamolies flagged out
    document.add_page_break()
    summary = document.add_heading('Image file Anomalies Summary', 0)
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    anomaly1 = overalldata[1]
    anomaly2 = overalldata[2]
    anomaly3 = overalldata[3]
    anomaly4 = overalldata[4]
    anomaly5 = overalldata[5]
    anomaly6 = overalldata[6]
    anomaly7 = overalldata[7]
    anomaly8 = overalldata[8]
    anomaly9 = overalldata[9]
    anomaly10 = overalldata[10]
    anomaly11 = overalldata[11]
    anomaly12 = overalldata[12]
    totalanomalynum = anomaly1 + anomaly2 + anomaly3 + anomaly4 + anomaly5 + anomaly6 + anomaly7 + anomaly8 + \
                      anomaly9 + anomaly10 + anomaly11 + anomaly12
    document.add_paragraph("The total number of anomalies being flagged out: " + str(totalanomalynum), style=None)
    anomaly_records = (
        (1, 'MFT record might be tampered/damaged', anomaly1),
        (2, '$SI Timestamp zerorised', anomaly2),
        (3, '$SI Timestamp is invalid', anomaly3),
        (4, '$SI Timestamp with 0 nanosecond', anomaly4),
        (5, '$SI Create Timestamp after $SI Modify Timestamp', anomaly5),
        (6, '$SI Assess Timestamp after both $SI Modify Timestamp and $SI Create Timestamp', anomaly6),
        (7, '$SI Timestamp after Current Time', anomaly7),
        (8, '$FN Timestamp zerorised', anomaly8),
        (9, '$FN Timestamp is invalid', anomaly9),
        (10, "$FN Timestamp with 0 nanosecond", anomaly10),
        (11, '$SI Create Timestamp after $FN Create Timestamp', anomaly11),
        (12, '$FN Timestamp after Current Time', anomaly12)
    )
    table = document.add_table(rows=1, cols=3, style='TableGrid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Anomaly Name'
    hdr_cells[2].text = 'Anomaly Count(s)'
    for num, anomal, anomaly_c in anomaly_records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = str(anomal)
        row_cells[2].text = str(anomaly_c)

# ************************************************ End of Anomalies Summary ****************************************** #
# *********************************************** Creation of Footer ************************************************* #
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(paragraph):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        page_run = paragraph.add_run()
        t1 = create_element('w:t')
        create_attribute(t1, 'xml:space', 'preserve')
        t1.text = 'Page '
        page_run._r.append(t1)

        page_num_run = paragraph.add_run()

        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)

        of_run = paragraph.add_run()
        t2 = create_element('w:t')
        create_attribute(t2, 'xml:space', 'preserve')
        t2.text = ' of '
        of_run._r.append(t2)

        fldChar3 = create_element('w:fldChar')
        create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = create_element('w:instrText')
        create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = create_element('w:fldChar')
        create_attribute(fldChar4, 'w:fldCharType', 'end')

        num_pages_run = paragraph.add_run()
        num_pages_run._r.append(fldChar3)
        num_pages_run._r.append(instrText2)
        num_pages_run._r.append(fldChar4)

    add_page_number(document.sections[0].footer.paragraphs[0])
    document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# ********************************************* End Creation of Footer *********************************************** #
    # Report Generation
    output_folder = './output/'
    document.save(output_folder + fileName + '.docx')
    print("\nGenerating PDF Document...")
    doc_report = output_folder + fileName + '.docx'
    pdf_report = output_folder + fileName + '.pdf'

    convert(doc_report)
    convert(doc_report, pdf_report)
    os.remove(doc_report)

# *********************************************** End Report Creation ************************************************ #

def args_route(mftfn):
    if(mftfn == None):
        print("Please input the MFT file.")
    else:
        f = open(mftfn, "rb")
        get_mft_eh_val(f)

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('-f',  help='-f <path to MFT file>')
    args = parser.parse_args()
    args_route(args.f)