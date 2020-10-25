import csv
import hashlib
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date

file_input = './input/out.csv'
sus_file = './input/test.sus'

class FileMetaExtractor:


    def __init__(self, file_path):
        self.file_path = file_path


    def get_file_records_count(self):
        try:
            return len(list(csv.reader(open(self.file_path))))
        except Exception as e:
            return -1
        return None

    # print("The total records in the MFT file image is: " + str(fileObject) + " records")


    def fileHash(self):
        try:
            with open(self.file_path, "rb") as f:  # rb stands opening the file in binary format for reading only
                hashed = f.read()  # read entire file as bytes
                return hashlib.sha256(hashed).hexdigest()
        except Exception as e:
            return None
        return None

    def fileSize(self):
        try:
            return os.path.getsize(self.file_path)
        except Exception as e:
            return -1
        # print("The size of the suspicious file is: " + str(getfileSize) + " bytes")


obj = FileMetaExtractor("./input/test.sus")
# print(obj.get_file_records_count())
# print(obj.fileHash())
# print(obj.fileSize())

document = Document()

header = document.sections[0].header
htable=header.add_table(1, 2, Inches(6))
htab_cells=htable.rows[0].cells
ht0=htab_cells[0].add_paragraph()
kh=ht0.add_run()
kh.add_picture('./images/SIT_logo.png', width=Inches(0.8))
ht1=htab_cells[1].add_paragraph('Singapore Institute of Technology')
ht1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

penis = document.add_heading('Digital Forensic Report', 0)
penis.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

reportCreationDate = document.add_paragraph("Report Creation Date: " + str(date.today()))
reportCreationDate.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

document.add_page_break()

penis1 = document.add_heading('Image File Summary', 0)
penis1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# document.add_paragraph('Company Address', style='Intense Quote')
# document.add_paragraph('Company Address', style='Intense Quote')
# document.add_paragraph('Company Address', style='Intense Quote')

# print(fileRecordsCount)

document.add_paragraph("The total records in the MFT file image is: " + str(obj.get_file_records_count()), style=None)
document.add_paragraph("The sha256sum of the suspicious file is: " + str(obj.fileHash()), style=None)
document.add_paragraph("The size of the suspicious file is: " + str(obj.fileSize()), style=None)


document.save('./output/poc.docx')